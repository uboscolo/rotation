"""
    This module provides utility functions for data processing.
    It includes functions for cleaning, transforming, and analyzing data.
"""
import argparse
from dataclasses import dataclass, field
import json
from pathlib import Path

rotate_map: dict = {
    3: {
        4: {1: [1,2,3], 2: [1,2,4], 3: [1,3,4], 4: [2,3,4],
            5: [1,2,3], 6: [1,2,4], 7: [1,3,4], 8: [2,3,4]},
        5: {1: [1,2,3], 2: [1,4,5], 3: [2,3,4], 4: [1,2,5],
            5: [3,4,5], 6: [1,2,4], 7: [1,3,5], 8: [2,4,5]},
        6: {1: [1,2,3], 2: [4,5,6], 3: [2,3,4], 4: [1,5,6],
            5: [3,4,5], 6: [1,2,6], 7: [1,3,5], 8: [2,4,6]},
    },
    5: {
        6: {1: [1,2,3,4,5], 2: [1,2,3,4,6], 3: [1,2,3,5,6], 4: [1,2,4,5,6],
            5: [1,3,4,5,6], 6: [2,3,4,5,6], 7: [1,2,3,4,5], 8: [1,2,3,4,6]},
        7: {1: [1,2,3,4,5], 2: [1,2,3,6,7], 3: [1,4,5,6,7], 4: [2,3,4,5,6],
            5: [1,2,3,4,7], 6: [1,2,5,6,7], 7: [3,4,5,6,7], 8: [1,2,3,4,5]},
        8: {1: [1,2,3,4,5], 2: [1,2,6,7,8], 3: [3,4,5,6,7], 4: [1,2,3,4,8],
            5: [1,5,6,7,8], 6: [2,3,4,5,6], 7: [1,2,3,7,8], 8: [4,5,6,7,8]},
    },
}

available_positions: dict = {
    "defense": {
        3: ["LB", "CB", "RB"],
    },
    "offense": {
        3: ["LW", "ST", "RW"],
        5: ["LM", "RM", "LW", "ST", "RW"],
    }

}

def load_positions(players_db: str) -> dict:
    """Loads player position preferences from the specified database file.
    
    Args:
        players_db (str): Path to the players database file.
    
    Returns:
        dict: The position preferences for all players.
    """
    positions_file = Path(players_db)
    return json.loads(positions_file.read_text(encoding="utf-8"))


class Player:
    """
    Represents a player
    """

    def __init__(self, name: str):
        """
        Initializes a new Player object.

        Args:
            name (str): The player's name.
        """
        self.__name: str = name
        self.__starting_position = None
        self.__position = None
        self.__favorite_position = None
        self.__alternate_position = None

    @property
    def name(self) -> str:
        """The name of the player."""
        return self.__name

    @name.setter
    def name(self, value: str):
        if not isinstance(value, str):
            raise TypeError("Name must be a string.")
        if not value:
            raise ValueError("Name cannot be empty.")
        self.__name = value

    @property
    def position(self) -> str:
        """The current position of the player."""
        return self.__position

    @position.setter
    def position(self, value: str):
        self.__position = value

    @property
    def starting_position(self) -> str:
        """The starting position of the player."""
        return self.__starting_position

    @starting_position.setter
    def starting_position(self, value: str):
        self.__starting_position = value

    @property
    def favorite_position(self) -> str:
        """The favorite position of the player."""
        return self.__favorite_position

    @favorite_position.setter
    def favorite_position(self, value: str):
        self.__favorite_position = value

    @property
    def alternate_position(self) -> str:
        """The alternate position of the player."""
        return self.__alternate_position

    @alternate_position.setter
    def alternate_position(self, value: str):
        self.__alternate_position = value

@dataclass
class UnitState:
    """Mutable state for a unit across periods."""
    starters: list = field(default_factory=list)
    reserves: list = field(default_factory=list)
    active_players: list = field(default_factory=list)
    inactive_players: list = field(default_factory=list)
    players_map: dict = field(default_factory=dict)
    assigned_positions: dict = field(default_factory=dict)

class Unit:
    """
    Represents a unit in a team, e.g. defense or offense
    """
    supported_starters = (3, 5)

    def __init__(self, name: str, num_starters: int):
        """
        Initializes a new Unit object.

        Args:
            name (str): The unit's name.
            num_starters (int): The number of starters in the unit.
        """
        if name not in available_positions:
            raise ValueError(f"Unit '{name}' is not a valid unit. "
                             f"Choose from: {list(available_positions.keys())}.")
        position_options = available_positions[name]
        if num_starters not in position_options:
            raise ValueError(f"Number of starters '{num_starters}' is not valid for unit '{name}'. "
                             f"Choose from: {list(position_options.keys())}.")
        self.__name: str = name
        self.__num_starters = num_starters
        self.__positions = position_options[num_starters]
        self.__state = UnitState()

    @property
    def num_starters(self) -> int:
        """The number of starters for this unit."""
        return self.__num_starters

    @property
    def name(self) -> str:
        """The name of the unit."""
        return self.__name

    @name.setter
    def name(self, value: str):
        if not isinstance(value, str):
            raise TypeError("Name must be a string.")
        if not value:
            raise ValueError("Name cannot be empty.")
        self.__name = value

    @property
    def starters(self) -> list:
        """The unit's starters."""
        return self.__state.starters

    @property
    def reserves(self) -> list:
        """The unit's reserves."""
        return self.__state.reserves

    @property
    def active_players(self) -> list:
        """The unit's active players."""
        return self.__state.active_players

    @property
    def inactive_players(self) -> list:
        """The unit's inactive players."""
        return self.__state.inactive_players

    def get_player_by_name(self, name: str) -> Player:
        """
        Returns a player given their name

        Args:
            name (str): The player's name.
        """
        return self.__state.players_map.get(name)

    def __set_position(self, player: Player):
        """
        Sets the position for the player based on their favorite and alternate positions.

        Args:
            player (Player): The player to set the position for.
        """
        favorite = player.favorite_position
        alternate = player.alternate_position
        position = None
        if favorite is not None and favorite not in self.__state.assigned_positions:
            position = favorite
        elif alternate is not None and alternate not in self.__state.assigned_positions:
            position = alternate
        else:
            # Neither preferred position is free; take first available slot
            for pos in self.__positions:
                if pos not in self.__state.assigned_positions:
                    position = pos
                    break
        if position is None:
            raise ValueError(f"No available position for player {player.name}.")
        self.__state.assigned_positions[position] = player.name
        return position

    def add_player(self, new_player: Player) -> None:
        """
        Adds a new player to the unit

        Args:
            new_player (Player): The player to add.
        """
        if new_player in self.__state.starters or new_player in self.__state.reserves:
            raise ValueError(f"Player {new_player.name} already in the {self.__name} unit.")

        current_starters = len(self.__state.starters)
        if current_starters == self.num_starters:
            # Reserves
            new_player.position = "-"
            new_player.starting_position = new_player.position
            self.reserves.append(new_player)
        else:
            new_player.position = self.__set_position(new_player)
            new_player.starting_position = new_player.position
            self.__state.starters.append(new_player)
        self.__state.players_map[new_player.name] = new_player
        # print(f"Adding {new_player.name} as {new_player.starting_position}")

    def swap(self, name: str, player: Player, **positions) -> None:
        """
        Swaps players

        Args:
            name (str): The player's name to swap out
            player (Player): The player to swap in
        """
        candidate = self.__state.players_map.get(name)
        default_starting_position = candidate.starting_position if candidate else None
        default_position = candidate.position if candidate else None
        starting_position = positions.get("starting_position", default_starting_position)
        position = positions.get("position", default_position)
        if not starting_position:
            raise ValueError("Starting position not defined.")
        if not position:
            raise ValueError("Active position not defined.")
        player.starting_position = starting_position
        player.position = position

        # Handle case of goalkeeper
        if starting_position == "-":
            self.__state.reserves[self.__state.reserves.index(candidate)] = player
        else:
            self.__state.starters[self.__state.starters.index(candidate)] = player
        if position == "-":
            self.__state.inactive_players[self.__state.inactive_players.index(candidate)] = player
        else:
            self.__state.active_players[self.__state.active_players.index(candidate)] = player
        self.__state.players_map.pop(name)
        self.__state.players_map[player.name] = player

    def __get_new_starters(self, period: int, size: int) -> list:
        """
         Returns the new starters given the unit size and the period

        Args:
            period (int): Period.
            size (int): Size.
        """
        new_starters = []
        if size > self.num_starters:
            rotate_by_size = rotate_map.get(self.num_starters, {})
            if not rotate_by_size:
                raise ValueError(
                    f"No rotation configuration for {self.num_starters} starters."
                )
            period_map = rotate_by_size.get(size, {})
            if not period_map:
                raise ValueError(
                    f"No rotation mapping for {self.num_starters} starters "
                    f"and unit size {size}."
                )

            # Rotation tables are defined for one cycle; reuse them for longer games.
            sorted_periods = sorted(period_map.keys())
            cycle_length = len(sorted_periods)
            normalized_period = ((period - 1) % cycle_length) + 1
            indices = period_map.get(normalized_period)
            if not indices:
                raise ValueError(
                    f"No rotation mapping for {self.num_starters} starters "
                    f"and unit size {size} at period {period}."
                )
            for i in indices:
                if i > len(self.__state.starters):
                    adjusted_index = i-1-self.num_starters
                    new_starter = self.__state.reserves[adjusted_index]
                else:
                    new_starter = self.__state.starters[i-1]
                new_starters.append(new_starter)
        else:
            new_starters = self.__state.starters.copy()
        return new_starters

    def __assign_with_alternative(self, target, player: Player,
                                  new_targets: dict,
                                  alternatives_available: dict) -> bool:
        """Assign player using an available alternative slot when target is occupied."""
        for alt_pos, alt_player in alternatives_available.items():
            if alt_pos not in new_targets:
                new_targets[alt_pos] = alt_player
                new_targets[target] = player
                return True
        return False

    def __assign_tuple_target(self, player: Player, target: tuple,
                              new_targets: dict,
                              alternatives_available: dict) -> bool:
        """Assign player for tuple targets, preferring first then second position."""
        pos1, pos2 = target
        if pos1 not in new_targets:
            new_targets[pos1] = player
            return True
        if pos2 not in new_targets:
            new_targets[pos2] = player
            return True
        return self.__assign_with_alternative(target, player,
                                              new_targets,
                                              alternatives_available)

    def __assign_single_target(self, player: Player, target: str,
                               new_targets: dict,
                               alternatives_available: dict) -> bool:
        """Assign player for single targets, falling back to alternative slots."""
        if target not in new_targets:
            new_targets[target] = player
            return True
        return self.__assign_with_alternative(target, player,
                                              new_targets,
                                              alternatives_available)

    def __collect_position_preferences(self) -> tuple:
        """Collect initial targets, alternatives, and swap candidates."""
        swappable = {}
        alternatives_available = {}
        new_targets = {}
        points = 0
        for player in self.__state.active_players:
            if player.favorite_position is None or player.alternate_position is None:
                continue
            if player.position == player.favorite_position:
                if player.favorite_position not in new_targets:
                    new_targets[player.favorite_position] = player
                if player.alternate_position not in alternatives_available:
                    alternatives_available[player.alternate_position] = player
                continue
            if player.position == player.alternate_position:
                points += 1
                swappable[player] = player.favorite_position
                continue
            points += 2
            swappable[player] = (player.favorite_position, player.alternate_position)
        return swappable, alternatives_available, new_targets, points

    def __reassign_swappables(self, swappable: dict,
                              alternatives_available: dict,
                              new_targets: dict) -> bool:
        """Try to assign each swappable player to a preferred target."""
        reassign = False
        for player, target in swappable.items():
            if isinstance(target, tuple):
                changed = self.__assign_tuple_target(player, target,
                                                     new_targets,
                                                     alternatives_available)
            else:
                changed = self.__assign_single_target(player, target,
                                                      new_targets,
                                                      alternatives_available)
            reassign = reassign or changed
        return reassign

    def __apply_targets(self, new_targets: dict) -> None:
        """Apply computed target positions to active players."""
        for player in self.__state.active_players:
            for new_pos, target_player in new_targets.items():
                if player == target_player:
                    player.position = new_pos

    def __validate_positions(self) -> None:
        """
        Rotates players

        Args:
            period (int): Period.
        """
        swappable, alternatives_available, new_targets, points = (
            self.__collect_position_preferences()
        )
        if points < len(self.__state.active_players) or not alternatives_available:
            return
        reassign = self.__reassign_swappables(swappable,
                                              alternatives_available,
                                              new_targets)
        if reassign:
            self.__apply_targets(new_targets)

    def rotate(self, period: int) -> None:
        """
        Rotates players

        Args:
            period (int): Period.
        """
        size = len(self.__state.starters) + len(self.__state.reserves)
        new_starters = self.__get_new_starters(period, size)
        if not self.__state.active_players:
            self.__state.active_players = new_starters
        else:
            for player in self.__state.active_players:
                if player not in new_starters:
                    self.__state.assigned_positions.pop(player.position)
                    player.position = "-"
                    self.__state.inactive_players.append(player)
        if not self.__state.inactive_players:
            self.__state.inactive_players = self.__state.reserves.copy()
        else:
            for player in new_starters:
                if player in self.__state.inactive_players:
                    player.position = self.__set_position(player)
                    self.__state.inactive_players.remove(player)
        self.__state.active_players = new_starters
        # for p in self.__active_players:
        #    print(f"1. {p.name} -> {p.position}")
        self.__validate_positions()
        # for p in self.__active_players:
        #    print(f"2. {p.name} -> {p.position}")

class Team:
    """
    Represents a team
    """

    supported_starters = (7, 9)

    def __init__(self, name: str, num_starters: int, positions_db: str):
        """
        Initializes a new Team object.

        Args:
            name (str): The team's name.
            num_starters (int): Number of starters in the team.
            positions_db (str): Path to the players database file.
        """
        if num_starters not in self.supported_starters:
            raise ValueError(f"Unsupported number of starters: {num_starters}. "
                             f"Supported values are: {self.supported_starters}.")

        if num_starters == 7:
            self.__defense: Unit = Unit("defense", num_starters=3)
            self.__offense: Unit = Unit("offense", num_starters=3)
        elif num_starters == 9:
            self.__defense: Unit = Unit("defense", num_starters=3)
            self.__offense: Unit = Unit("offense", num_starters=5)

        self.__name: str = name
        self.__players = []
        self.__goalkeeper = None
        self.__goalkeeper_reserve = []
        self.__positions_db = load_positions(positions_db)
        self.__players_map = {}

    @property
    def name(self) -> str:
        """The name of the team."""
        return self.__name

    @name.setter
    def name(self, value: str):
        if not isinstance(value, str):
            raise TypeError("Name must be a string.")
        if not value:
            raise ValueError("Name cannot be empty.")
        self.__name = value

    @property
    def players(self) -> list:
        """The list of players."""
        return self.__players

    def get_active_players(self) -> list:
        """The list of active players."""
        active_players = [self.__goalkeeper]
        active_players.extend(self.__defense.active_players)
        active_players.extend(self.__offense.active_players)
        return active_players

    def get_players(self) -> list:
        """The list of all players."""
        all_players = [self.__goalkeeper]
        all_players.extend(self.__defense.active_players)
        all_players.extend(self.__defense.inactive_players)
        all_players.extend(self.__offense.active_players)
        all_players.extend(self.__offense.inactive_players)
        return all_players


    def get_player_by_name(self, name: str) -> Player:
        """
        Returns a player given their name

        Args:
            name (str): The player's name.
        """
        return self.__players_map.get(name)

    def add_goalkeeper(self, player: Player) -> None:
        """
        Adds goalkeepers to the team

        Args:
            player (Player): The player to add.
        """
        if not self.__goalkeeper:
            if self.get_player_by_name(player.name):
                raise ValueError(f"Player {player.name} already added.")
            player.starting_position = "GK"
            player.position = "GK"
            self.__goalkeeper = player
            self.__players_map[player.name] = player
            # print(f"Adding {player.name} as {player.position}")
        else:
            if player not in self.__players and player != self.__goalkeeper:
                raise ValueError(f"Player {player.name} not part of the team.")
            self.__goalkeeper_reserve.append(player)

    def add_player(self, new_player: Player, unit_name: str) -> None:
        """
        Adds a new player to the team

        Args:
            new_player (Player): The player to add.
            unit_name (str): The unit's name where to add the player.
        """
        if new_player in self.__players:
            raise ValueError(f"Player {new_player.name} already in the team.")

        if not self.__positions_db.get(new_player.name):
            raise ValueError(f"Player {new_player.name} not found in the database.")
        player_units = self.__positions_db[new_player.name]

        if not player_units.get(unit_name):
            raise ValueError(f"Player {new_player.name} not found in the {unit_name} unit.")
        player_unit = player_units[unit_name]

        favorite = player_unit.get("favorite")
        if favorite is None:
            raise ValueError(f"Player {new_player.name} has no favorite position in the {unit_name} unit.")
        new_player.favorite_position = favorite

        alternate = player_unit.get("alternate")
        if alternate is None:
            raise ValueError(f"Player {new_player.name} has no alternate position in the {unit_name} unit.")
        new_player.alternate_position = alternate

        self.__players.append(new_player)
        self.__players_map[new_player.name] = new_player
        unit_map = {"defense": self.__defense, "offense": self.__offense}
        unit = unit_map.get(unit_name)
        unit.add_player(new_player)

    def __update_preferences(self, player: Player, unit_name: str) -> None:
        """Update a player's favorite/alternate to match their new unit."""
        player_units = self.__positions_db.get(player.name, {})
        unit_prefs = player_units.get(unit_name, {})
        if unit_prefs.get("favorite"):
            player.favorite_position = unit_prefs["favorite"]
        if unit_prefs.get("alternate"):
            player.alternate_position = unit_prefs["alternate"]

    def swap(self, player_name1: str, player_name2: str) -> None:
        """
        Swaps players

        Args:
            player_name1 (str): Player to swap.
            player_name2 (str): Player to swap.
        """
        if self.__defense.get_player_by_name(player_name1):
            player1 = self.__defense.get_player_by_name(player_name1)
            player2 = self.__offense.get_player_by_name(player_name2)
            if not player2:
                if player_name2 != self.__goalkeeper.name:
                    raise ValueError(f"Player2 {player_name2} not in offense.")
                self.__defense.swap(player_name1, self.__goalkeeper)
                self.__swap_keeper(player1, True)
            else:
                positions = {"starting_position": player2.starting_position,
                             "position": player2.position}
                self.__defense.swap(player_name1, player2)
                self.__offense.swap(player_name2, player1, **positions)
                self.__update_preferences(player2, "defense")
                self.__update_preferences(player1, "offense")
        elif self.__offense.get_player_by_name(player_name1):
            player1 = self.__offense.get_player_by_name(player_name1)
            player2 = self.__defense.get_player_by_name(player_name2)
            if not player2:
                if player_name2 != self.__goalkeeper.name:
                    raise ValueError(f"Player2 {player_name2} not in defense.")
                self.__offense.swap(player_name1, self.__goalkeeper)
                self.__swap_keeper(player1, True)
            else:
                positions = {"starting_position": player2.starting_position,
                             "position": player2.position}
                self.__offense.swap(player_name1, player2)
                self.__defense.swap(player_name2, player1, **positions)
                self.__update_preferences(player1, "defense")
                self.__update_preferences(player2, "offense")
        else:
            # could it be goalkeeper
            if player_name1 != self.__goalkeeper.name:
                raise ValueError(f"Player to swap not found: {player_name1}.")
            player2 = self.__defense.get_player_by_name(player_name2)
            self.__swap_keeper(player2)

    def __swap_keeper(self, candidate: Player, force: bool = False) -> None:
        """
        Swaps goalkeeper

        Args:
            candidate (Player): The player to be swapped in.
            force (bool): Forces the swap even if candidate not in any unit.
        """
        if candidate != self.__goalkeeper:
            ex_keeper = self.__goalkeeper
            if self.__defense.get_player_by_name(candidate.name):
                self.__defense.swap(candidate.name, ex_keeper)
                self.__update_preferences(ex_keeper, "defense")
            elif self.__offense.get_player_by_name(candidate.name):
                self.__offense.swap(candidate.name, ex_keeper)
                self.__update_preferences(ex_keeper, "offense")
            else:
                if not force:
                    raise ValueError(f"Candidate {candidate.name} not in team.")

            candidate.position = "GK"
            self.__goalkeeper = candidate

    def rotate(self, period: int) -> None:
        """
        Rotates players

        Args:
            period (int): Period.
        """
        if period < 1 or period > 8:
            raise ValueError(f"Unexpected period {period}.")
        if period > 1 and period % 2 == 1:
            candidate = self.__goalkeeper_reserve.pop(0)
            self.__swap_keeper(candidate)
        self.__defense.rotate(period)
        self.__offense.rotate(period)

class Period:
    """
    Represents a period in a game
    """
    def __init__(self, number: int) -> None:
        """
        Initializes a new Period object.

        Args:
            number (int): The period's number.
        """
        self.__number = number
        self.__positions = {}

    @property
    def number(self) -> int:
        """The number of the period."""
        return self.__number

    @property
    def positions(self) -> dict:
        """The positions of the period."""
        return self.__positions

    def add_position(self, player_name: str, position: str) -> None:
        """
        Adds a player's position to the period

        Args:
            player_name (str): The name of the player.
            position (str): The position of the player.
        """
        self.__positions[player_name] = position

class Half:
    """
    Represents a half in a game
    """
    def __init__(self, number: int) -> None:
        """
        Initializes a new Half object.

        Args:
            number (int): The half's number.
        """
        self.__number = number
        self.__periods = []

    @property
    def number(self) -> int:
        """The number of the half."""
        return self.__number

    def add_period(self, period: Period) -> None:
        """
        Adds a period to the half

        Args:
            period (Period): The period to add.
        """
        self.__periods.append(period)

    def get_table_data(self) -> tuple:
        """Build header labels and per-player position rows for this half."""
        positions = {}
        headers = []
        for period in self.__periods:
            for player_name, position in period.positions.items():
                if player_name not in positions:
                    positions[player_name] = [position]
                else:
                    positions[player_name].append(position)
            if period.number % 4 == 1:
                headers.append("Starting")
            elif period.number % 4 == 2:
                headers.append("1st sub")
            elif period.number % 4 == 3:
                headers.append("2nd sub")
            else:
                headers.append("3rd sub")
        return headers, positions

    def display(self) -> None:
        """
        Displays the half information.
        """
        print(f"Half {self.__number}")
        print(f"{'':<12}", end="")
        headers, positions = self.get_table_data()
        for i, header in enumerate(headers):
            line_end = "" if i < len(headers) - 1 else "\n"
            print(f"{header:<12}", end=line_end)
        for player_name, pos_list in positions.items():
            print(f"{player_name:<12}", end="")
            for pos in pos_list:
                print(f"{pos:<12}", end="")
            print("")

class Game:
    """
    Represents a game
    """
    def __init__(self) -> None:
        """
        Initializes a new Game object.
        """
        self.__halves = []

    def add_half(self, half: Half) -> None:
        """
        Adds a half to the game

        Args:
            half (Half): The half to add.
        """
        self.__halves.append(half)

    def display(self) -> None:
        """
        Displays the game information.
        """
        for half in self.__halves:
            half.display()

    def write_docx(self, output_file: str) -> None:
        """Write game rotation tables to a DOCX file."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX export. Install it with: pip install python-docx"
            ) from exc

        document = Document()

        def _set_cell_font_size(cell, size_pt: int, bold: bool = False) -> None:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)
                    run.font.bold = bold

        for half_idx, half in enumerate(self.__halves):
            half_label = "1st Half" if half.number == 1 else "2nd Half"
            headers, positions = half.get_table_data()
            table = document.add_table(rows=len(positions) + 2, cols=len(headers) + 1)
            table.style = "Table Grid"
            title_cell = table.cell(0, 0)
            title_cell.text = half_label
            for col in range(1, len(headers) + 1):
                title_cell = title_cell.merge(table.cell(0, col))
            for paragraph in table.cell(0, 0).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(24)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

            table.cell(1, 0).text = "Player"
            _set_cell_font_size(table.cell(1, 0), 18, bold=True)
            for i, header in enumerate(headers, start=1):
                table.cell(1, i).text = header
                _set_cell_font_size(table.cell(1, i), 18, bold=True)
            for row_idx, (player_name, pos_list) in enumerate(positions.items(), start=2):
                table.cell(row_idx, 0).text = player_name
                _set_cell_font_size(table.cell(row_idx, 0), 18)
                for col_idx, pos in enumerate(pos_list, start=1):
                    table.cell(row_idx, col_idx).text = str(pos)
                    _set_cell_font_size(table.cell(row_idx, col_idx), 18)
            # Add page break after each half except the last one
            if half_idx < len(self.__halves) - 1:
                document.add_page_break()
            else:
                document.add_paragraph("")
        document.save(output_file)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Rotation")
    parser.add_argument("--players_file",
                        type=str,
                        required=True,
                        dest="players_file",
                        help="Path to players file")
    parser.add_argument("--players_db",
                        type=str,
                        required=False,
                        dest="players_db",
                        default="database/players.json",
                        help="Path to players database file")
    parser.add_argument("--num_starters",
                        type=int,
                        required=True,
                        dest="num_starters",
                        help="Number of starters in the game")
    parser.add_argument("--num_periods",
                        type=int,
                        required=False,
                        dest="num_periods",
                        default=8,
                        help="Number of periods in the game")
    parser.add_argument("--docx_file",
                        type=str,
                        required=False,
                        dest="docx_file",
                        default=None,
                        help="Optional path to write a DOCX rotation report")
    args = parser.parse_args()
    in_file = args.players_file
    with open(in_file, "r", encoding="utf-8") as f:
        data = json.load(f)
        defense = data.get("defense", [])
        offense = data.get("offense", [])
        keepers = data.get("keepers", [])
        half_time_swaps = data.get("swaps", [])

    num_periods = args.num_periods
    team = Team("B4 Lions", args.num_starters, args.players_db)
    for defender_name in defense:
        defender_player = Player(defender_name)
        team.add_player(defender_player, "defense")
    for attacker_name in offense:
        attacker_player = Player(attacker_name)
        team.add_player(attacker_player, "offense")
    starting_keeper_name = keepers.pop(0)
    starting_keeper = Player(starting_keeper_name)
    team.add_goalkeeper(starting_keeper)
    for reserve_keeper_name in keepers:
        reserve_keeper = team.get_player_by_name(reserve_keeper_name)
        team.add_goalkeeper(reserve_keeper)

    new_game = Game()
    for period_number in range(1, num_periods+1):
        if period_number == 1:
            new_half = Half(1)
            new_game.add_half(new_half)
        if period_number == 5:
            new_half = Half(2)
            new_game.add_half(new_half)
            for swap in half_time_swaps:
                team.swap(swap[0], swap[1])
        team.rotate(period_number)
        new_period = Period(period_number)
        for team_player in team.get_players():
            new_period.add_position(team_player.name, team_player.position)
        new_half.add_period(new_period)
    new_game.display()
    if args.docx_file:
        new_game.write_docx(args.docx_file)
        print(f"DOCX report written to {args.docx_file}")
